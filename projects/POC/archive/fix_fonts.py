#!/usr/bin/env python3
"""
fix_fonts.py — Standardize body-text fonts in agentic-cfa-spec.docx to
Times New Roman 10pt, leaving headings/titles/captions/TOC untouched.

Pre-mortem mitigations:
  Risk 1: File path confirmed before any write (backup before modify).
  Risk 2: All paragraph styles logged before changes; heuristic reviewed.
  Risk 3: Timestamped backup created; file sizes compared post-save.
  Risk 4: Inline run-level overrides explicitly set on every run.
  Risk 5: Table cells, headers, and footers handled explicitly.
  Risk 6: Style list printed for manual review of edge cases.
"""

import shutil
import sys
import os
from pathlib import Path
from datetime import datetime

# ── Dependency check ──────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
except ImportError:
    print("[setup] python-docx not found — installing...", flush=True)
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn

# ── Constants ─────────────────────────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).parent
TARGET_FILE = SCRIPT_DIR / "agentic-cfa-spec.docx"
BACKUP_FILE = SCRIPT_DIR / "agentic-cfa-spec.backup.docx"
TARGET_FONT = "Times New Roman"
TARGET_SIZE = Pt(10)

# Style name prefixes that mark non-body content — do NOT touch these.
EXCLUDED_PREFIXES = ("Heading", "Title", "Subtitle", "Caption", "TOC")


def is_body_style(style_name: str) -> bool:
    """Return True if the style name indicates body text (not a heading/title/etc.)."""
    if style_name is None:
        return True
    for prefix in EXCLUDED_PREFIXES:
        if style_name.startswith(prefix):
            return False
    return True


def style_name_of(para) -> str:
    """Return the paragraph's style name, defaulting to 'Normal'."""
    try:
        return para.style.name if para.style and para.style.name else "Normal"
    except Exception:
        return "Normal"


def apply_font_to_paragraph(para, changed_paras: list, changed_runs: list):
    """Apply Times New Roman 10pt to every run in a body-text paragraph."""
    sname = style_name_of(para)
    if not is_body_style(sname):
        return
    changed_paras.append(sname)
    for run in para.runs:
        run.font.name = TARGET_FONT
        run.font.size = TARGET_SIZE
        changed_runs.append((sname, run.text[:40] if run.text else ""))


# ── Step 1: Verify file exists ────────────────────────────────────────────────
print("=" * 70)
print("STEP 1: File verification")
print("=" * 70)
if not TARGET_FILE.exists():
    print(f"[ERROR] File not found: {TARGET_FILE}", flush=True)
    sys.exit(1)
original_size = TARGET_FILE.stat().st_size
print(f"  Target  : {TARGET_FILE}")
print(f"  Size    : {original_size:,} bytes", flush=True)

# ── Step 2: Backup ────────────────────────────────────────────────────────────
print()
print("=" * 70)
print("STEP 2: Backup")
print("=" * 70)
shutil.copy2(TARGET_FILE, BACKUP_FILE)
if not BACKUP_FILE.exists():
    print(f"[ERROR] Backup failed — aborting.", flush=True)
    sys.exit(1)
backup_size = BACKUP_FILE.stat().st_size
print(f"  Backup  : {BACKUP_FILE}")
print(f"  Size    : {backup_size:,} bytes")
print(f"  Timestamp: {datetime.now().isoformat()}", flush=True)

# ── Step 3: Open and inspect document ─────────────────────────────────────────
print()
print("=" * 70)
print("STEP 3: Document inspection (pre-change)")
print("=" * 70)

doc = Document(TARGET_FILE)

# 3a: Collect all unique style names across body paragraphs and table cells
all_style_names: set = set()
for para in doc.paragraphs:
    all_style_names.add(style_name_of(para))
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                all_style_names.add(style_name_of(para))

print(f"\n  Paragraph styles found in document:")
body_styles = sorted(s for s in all_style_names if is_body_style(s))
excluded_styles = sorted(s for s in all_style_names if not is_body_style(s))
print(f"    Body-text styles (will be changed):")
for s in body_styles:
    print(f"      - {s}")
print(f"    Excluded styles (headings/titles/etc., will NOT be changed):")
for s in excluded_styles:
    print(f"      - {s}")

# 3b: Table inventory
table_count = len(doc.tables)
row_count = sum(len(t.rows) for t in doc.tables)
cell_count = sum(len(r.cells) for t in doc.tables for r in t.rows)
print(f"\n  Tables : {table_count}")
print(f"  Rows   : {row_count}")
print(f"  Cells  : {cell_count}")

# 3c: Headers/footers
sections = list(doc.sections)
print(f"\n  Sections: {len(sections)}")
for i, section in enumerate(sections):
    try:
        hdr_text = " | ".join(p.text for p in section.header.paragraphs if p.text.strip())
        ftr_text = " | ".join(p.text for p in section.footer.paragraphs if p.text.strip())
        print(f"    Section {i+1} header: {hdr_text!r}")
        print(f"    Section {i+1} footer: {ftr_text!r}")
    except Exception as e:
        print(f"    Section {i+1}: could not read header/footer ({e})")

# 3d: Inline font overrides (runs where font.name is explicitly set)
inline_overrides = 0
for para in doc.paragraphs:
    for run in para.runs:
        if run.font.name is not None:
            inline_overrides += 1
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.font.name is not None:
                        inline_overrides += 1
print(f"\n  Runs with inline font overrides: {inline_overrides}")

# 3e: Footnotes check
has_footnotes = False
try:
    footnotes_part = doc.part.package.part_related_by(
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
    )
    has_footnotes = True
    print(f"\n  Footnotes: DETECTED (not modified — review manually if needed)")
except Exception:
    try:
        # Alternative check via XML
        pkg = doc.part.package
        rels = pkg._rels
        for rel in rels.values() if hasattr(rels, 'values') else []:
            if 'footnotes' in getattr(rel, 'reltype', ''):
                has_footnotes = True
                break
    except Exception:
        pass
    if not has_footnotes:
        print(f"\n  Footnotes: none detected")

# ── Step 4: Apply font changes ────────────────────────────────────────────────
print()
print("=" * 70)
print("STEP 4: Applying font changes")
print("=" * 70)

changed_paras: list = []
changed_runs: list = []

# 4a: Update document-level Normal style
try:
    normal_style = doc.styles["Normal"]
    normal_style.font.name = TARGET_FONT
    normal_style.font.size = TARGET_SIZE
    print(f"  [OK] Updated 'Normal' document style")
except KeyError:
    print(f"  [WARN] 'Normal' style not found — skipping document-level update")

# 4b: Main body paragraphs
body_para_count = 0
for para in doc.paragraphs:
    if is_body_style(style_name_of(para)):
        apply_font_to_paragraph(para, changed_paras, changed_runs)
        body_para_count += 1
print(f"  Main body: {body_para_count} body paragraphs processed")

# 4c: Table cells
table_para_count = 0
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if is_body_style(style_name_of(para)):
                    apply_font_to_paragraph(para, changed_paras, changed_runs)
                    table_para_count += 1
print(f"  Tables  : {table_para_count} body paragraphs processed")

# 4d: Headers and footers
hf_para_count = 0
for section in doc.sections:
    try:
        for para in section.header.paragraphs:
            if is_body_style(style_name_of(para)):
                apply_font_to_paragraph(para, changed_paras, changed_runs)
                hf_para_count += 1
        for para in section.footer.paragraphs:
            if is_body_style(style_name_of(para)):
                apply_font_to_paragraph(para, changed_paras, changed_runs)
                hf_para_count += 1
    except Exception as e:
        print(f"  [WARN] Header/footer section error: {e}")
print(f"  Headers/footers: {hf_para_count} body paragraphs processed")

total_paras = body_para_count + table_para_count + hf_para_count
print(f"\n  Total paragraphs changed: {total_paras}")
print(f"  Total runs changed      : {len(changed_runs)}", flush=True)

# ── Step 5: Save in-place ─────────────────────────────────────────────────────
print()
print("=" * 70)
print("STEP 5: Saving")
print("=" * 70)
doc.save(TARGET_FILE)
saved_size = TARGET_FILE.stat().st_size
print(f"  Saved to: {TARGET_FILE}")
print(f"  Size    : {saved_size:,} bytes")
size_delta_pct = ((saved_size - original_size) / original_size) * 100 if original_size else 0
print(f"  Delta vs original: {size_delta_pct:+.1f}%")
if saved_size < original_size * 0.90:
    print(f"  [WARNING] Saved file is >10% smaller than original — possible truncation!", flush=True)
else:
    print(f"  [OK] File size looks normal", flush=True)

# ── Step 6: Validate ──────────────────────────────────────────────────────────
print()
print("=" * 70)
print("STEP 6: Validation")
print("=" * 70)
try:
    doc_check = Document(TARGET_FILE)
    para_count_check = len(doc_check.paragraphs)
    print(f"  [OK] Re-opened successfully ({para_count_check} top-level paragraphs)")
except Exception as e:
    print(f"  [ERROR] Failed to re-open saved document: {e}", flush=True)
    sys.exit(1)

# ── Step 7: Post-scan exception report ───────────────────────────────────────
print()
print("=" * 70)
print("STEP 7: Post-scan exception report")
print("=" * 70)

doc_final = Document(TARGET_FILE)
exception_font: list = []
exception_size: list = []

def scan_para_exceptions(para, location: str):
    sname = style_name_of(para)
    if not is_body_style(sname):
        return
    for run in para.runs:
        if run.font.name is not None and run.font.name != TARGET_FONT:
            exception_font.append((location, sname, run.font.name, run.text[:60]))
        if run.font.size is not None and run.font.size != TARGET_SIZE:
            exception_size.append((location, sname, str(run.font.size), run.text[:60]))

for para in doc_final.paragraphs:
    scan_para_exceptions(para, "body")
for table in doc_final.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                scan_para_exceptions(para, "table")
for section in doc_final.sections:
    try:
        for para in section.header.paragraphs:
            scan_para_exceptions(para, "header")
        for para in section.footer.paragraphs:
            scan_para_exceptions(para, "footer")
    except Exception:
        pass

if exception_font:
    print(f"\n  FONT exceptions ({len(exception_font)} runs still NOT Times New Roman):")
    for loc, sname, fname, text in exception_font[:20]:
        print(f"    [{loc}] style={sname!r} font={fname!r} text={text!r}")
    if len(exception_font) > 20:
        print(f"    ... and {len(exception_font)-20} more")
else:
    print(f"\n  Font exceptions: 0 — all body runs are Times New Roman ✓")

if exception_size:
    print(f"\n  SIZE exceptions ({len(exception_size)} runs still NOT 10pt):")
    for loc, sname, sz, text in exception_size[:20]:
        print(f"    [{loc}] style={sname!r} size={sz} text={text!r}")
    if len(exception_size) > 20:
        print(f"    ... and {len(exception_size)-20} more")
else:
    print(f"\n  Size exceptions: 0 — all body runs are 10pt ✓")

# ── Final summary ─────────────────────────────────────────────────────────────
print()
print("=" * 70)
print("SUMMARY")
print("=" * 70)
print(f"  Backup     : {BACKUP_FILE} ({backup_size:,} bytes)")
print(f"  Output     : {TARGET_FILE} ({saved_size:,} bytes)")
print(f"  Paragraphs changed : {total_paras}")
print(f"  Runs changed       : {len(changed_runs)}")
print(f"  Font exceptions    : {len(exception_font)}")
print(f"  Size exceptions    : {len(exception_size)}")
print(f"  Footnotes detected : {'yes (review manually)' if has_footnotes else 'no'}")
print(f"  Validation         : OK")
print()
if len(exception_font) == 0 and len(exception_size) == 0:
    print("  ✓ SUCCESS — document body text fully standardized to Times New Roman 10pt")
else:
    print("  ⚠  Exceptions remain — review the exception report above and decide")
    print("     whether to extend the fix or handle those cases manually.")
print("=" * 70, flush=True)
